import io
import os
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.http import HttpResponse
from django.contrib import messages
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import (
    Projeto, PastaDeTrabalho, RegistroDesenvolvimento, FonteDeDados, FiltroFonte,
    ImagemFonte, Planilha, ImagemPlanilha, Painel, ImagemPainel,
    Historia, ImagemHistoria,
)
from .forms import (
    ProjetoForm, PastaDeTrabalhoForm, RegistroDesenvolvimentoForm,
    FonteDeDadosForm, FiltroFonteForm,
    ImagemFonteForm, PlanilhaForm, ImagemPlanilhaForm, PainelForm,
    ImagemPainelForm, HistoriaForm, ImagemHistoriaForm,
)


def _url_projeto(projeto_pk, aba='dados'):
    return reverse('detalhe_projeto', args=[projeto_pk]) + f'?aba={aba}'


# ── Projetos ──────────────────────────────────────────────────────────────────

def lista_projetos(request):
    projetos = Projeto.objects.prefetch_related('pastas').all()
    return render(request, 'paineis/lista_projetos.html', {'projetos': projetos})


def criar_projeto(request):
    form = ProjetoForm(request.POST or None)
    if form.is_valid():
        projeto = form.save()
        messages.success(request, 'Projeto criado.')
        return redirect(_url_projeto(projeto.pk, 'projeto'))
    return render(request, 'paineis/form_projeto.html', {'form': form, 'titulo': 'Novo Projeto'})


def excluir_projeto(request, pk):
    projeto = get_object_or_404(Projeto, pk=pk)
    if request.method == 'POST':
        projeto.delete()
        messages.success(request, 'Projeto excluído.')
        return redirect('lista_projetos')
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': projeto,
        'tipo': 'projeto',
        'voltar_url': 'detalhe_projeto',
        'voltar_pk': projeto.pk,
    })


def detalhe_projeto(request, pk):
    """Página principal — abas: Projeto, Dados do Tableau, Fonte de Dados, Planilha, Painel, História."""
    projeto = get_object_or_404(Projeto, pk=pk)
    pasta = projeto.pastas.first()
    aba = request.GET.get('aba', 'projeto')

    # Bloqueia abas de conteúdo enquanto Dados do Tableau não estiver salvo
    if not pasta and aba not in ('projeto', 'dados'):
        return redirect(_url_projeto(pk, 'dados'))

    projeto_form = ProjetoForm(instance=projeto)
    form = PastaDeTrabalhoForm(instance=pasta) if pasta else PastaDeTrabalhoForm(initial={'nome': projeto.nome})
    reg_form = RegistroDesenvolvimentoForm()

    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        if form_type == 'projeto':
            projeto_form = ProjetoForm(request.POST, instance=projeto)
            if projeto_form.is_valid():
                projeto_form.save()
                messages.success(request, 'Projeto atualizado.')
                return redirect(_url_projeto(pk, 'projeto'))

        elif form_type == 'dados':
            if pasta:
                form = PastaDeTrabalhoForm(request.POST, instance=pasta)
                if form.is_valid():
                    form.save()
                    messages.success(request, 'Dados do Tableau atualizados.')
                    return redirect(_url_projeto(pk, 'dados'))
            else:
                form = PastaDeTrabalhoForm(request.POST)
                if form.is_valid():
                    nova_pasta = form.save(commit=False)
                    nova_pasta.projeto = projeto
                    nova_pasta.save()
                    messages.success(request, 'Dados do Tableau cadastrados.')
                    return redirect(_url_projeto(pk, 'dados'))

    ctx = {
        'projeto': projeto,
        'pasta': pasta,
        'aba': aba,
        'projeto_form': projeto_form,
        'form': form,
        'reg_form': reg_form,
    }

    if pasta:
        ctx.update({
            'registros': pasta.registros.all(),
            'fontes': pasta.fontes.prefetch_related('filtros', 'imagens').all(),
            'planilhas': pasta.planilhas.prefetch_related('imagens').all(),
            'paineis_list': pasta.paineis.prefetch_related('imagens', 'planilhas').all(),
            'historias': pasta.historias.prefetch_related('imagens').all(),
        })

    return render(request, 'paineis/detalhe_projeto.html', ctx)


# ── Registro de Desenvolvimento ───────────────────────────────────────────────

def adicionar_registro(request, pasta_pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pasta_pk)
    form = RegistroDesenvolvimentoForm(request.POST or None)
    if form.is_valid():
        registro = form.save(commit=False)
        registro.pasta = pasta
        registro.tipo = 'atualizado'
        registro.save()
        messages.success(request, 'Registro de atualização adicionado.')
        return redirect(_url_projeto(pasta.projeto.pk, 'dados'))
    return render(request, 'paineis/form_registro.html', {
        'form': form, 'pasta': pasta, 'projeto': pasta.projeto,
    })


def excluir_registro(request, pk):
    registro = get_object_or_404(RegistroDesenvolvimento, pk=pk)
    projeto_pk = registro.pasta.projeto.pk
    if request.method == 'POST':
        registro.delete()
        return redirect(_url_projeto(projeto_pk, 'dados'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': registro,
        'tipo': 'registro',
        'voltar_url': 'detalhe_projeto',
        'voltar_pk': projeto_pk,
    })


def excluir_pasta(request, pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pk)
    projeto_pk = pasta.projeto.pk
    if request.method == 'POST':
        pasta.delete()
        messages.success(request, 'Dados do Tableau excluídos.')
        return redirect(_url_projeto(projeto_pk, 'dados'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': pasta,
        'tipo': 'dados do Tableau',
        'voltar_url': 'detalhe_projeto',
        'voltar_pk': projeto_pk,
    })


# ── Fontes de Dados ───────────────────────────────────────────────────────────

def criar_fonte(request, pasta_pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pasta_pk)
    projeto = pasta.projeto
    form = FonteDeDadosForm(request.POST or None)
    if form.is_valid():
        fonte = form.save(commit=False)
        fonte.pasta = pasta
        fonte.save()
        messages.success(request, 'Fonte de dados criada.')
        return redirect('detalhe_fonte', pk=fonte.pk)
    return render(request, 'paineis/form_fonte.html', {
        'form': form, 'pasta': pasta, 'projeto': projeto, 'titulo': 'Nova Fonte de Dados',
    })


def detalhe_fonte(request, pk):
    fonte = get_object_or_404(FonteDeDados, pk=pk)
    return render(request, 'paineis/detalhe_fonte.html', {
        'fonte': fonte, 'projeto': fonte.pasta.projeto,
    })


def editar_fonte(request, pk):
    fonte = get_object_or_404(FonteDeDados, pk=pk)
    form = FonteDeDadosForm(request.POST or None, instance=fonte)
    if form.is_valid():
        form.save()
        messages.success(request, 'Fonte de dados atualizada.')
        return redirect('detalhe_fonte', pk=fonte.pk)
    return render(request, 'paineis/form_fonte.html', {
        'form': form, 'pasta': fonte.pasta, 'projeto': fonte.pasta.projeto,
        'titulo': 'Editar Fonte de Dados',
    })


def excluir_fonte(request, pk):
    fonte = get_object_or_404(FonteDeDados, pk=pk)
    projeto_pk = fonte.pasta.projeto.pk
    if request.method == 'POST':
        fonte.delete()
        messages.success(request, 'Fonte de dados excluída.')
        return redirect(_url_projeto(projeto_pk, 'fontes'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': fonte,
        'tipo': 'fonte de dados',
        'voltar_url': 'detalhe_fonte',
        'voltar_pk': fonte.pk,
    })


def criar_filtro(request, fonte_pk):
    fonte = get_object_or_404(FonteDeDados, pk=fonte_pk)
    form = FiltroFonteForm(request.POST or None)
    if form.is_valid():
        filtro = form.save(commit=False)
        filtro.fonte = fonte
        filtro.save()
        messages.success(request, 'Filtro adicionado.')
        return redirect('detalhe_fonte', pk=fonte.pk)
    return render(request, 'paineis/form_filtro.html', {'form': form, 'fonte': fonte})


def editar_filtro(request, pk):
    filtro = get_object_or_404(FiltroFonte, pk=pk)
    form = FiltroFonteForm(request.POST or None, instance=filtro)
    if form.is_valid():
        form.save()
        messages.success(request, 'Filtro atualizado.')
        return redirect('detalhe_fonte', pk=filtro.fonte.pk)
    return render(request, 'paineis/form_filtro.html', {'form': form, 'fonte': filtro.fonte})


def excluir_filtro(request, pk):
    filtro = get_object_or_404(FiltroFonte, pk=pk)
    fonte_pk = filtro.fonte.pk
    if request.method == 'POST':
        filtro.delete()
        messages.success(request, 'Filtro excluído.')
        return redirect('detalhe_fonte', pk=fonte_pk)
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': filtro,
        'tipo': 'filtro',
        'voltar_url': 'detalhe_fonte',
        'voltar_pk': fonte_pk,
    })


def adicionar_imagem_fonte(request, fonte_pk):
    fonte = get_object_or_404(FonteDeDados, pk=fonte_pk)
    form = ImagemFonteForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        imagem = form.save(commit=False)
        imagem.fonte = fonte
        imagem.save()
        messages.success(request, 'Imagem adicionada.')
        return redirect('detalhe_fonte', pk=fonte.pk)
    return render(request, 'paineis/form_imagem.html', {
        'form': form, 'titulo': 'Adicionar Imagem',
        'voltar_url': 'detalhe_fonte', 'voltar_pk': fonte.pk,
    })


def excluir_imagem_fonte(request, pk):
    imagem = get_object_or_404(ImagemFonte, pk=pk)
    fonte_pk = imagem.fonte.pk
    if request.method == 'POST':
        if imagem.imagem and os.path.isfile(imagem.imagem.path):
            os.remove(imagem.imagem.path)
        imagem.delete()
        messages.success(request, 'Imagem excluída.')
        return redirect('detalhe_fonte', pk=fonte_pk)
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': imagem, 'tipo': 'imagem',
        'voltar_url': 'detalhe_fonte', 'voltar_pk': fonte_pk,
    })


# ── Planilhas ─────────────────────────────────────────────────────────────────

def criar_planilha(request, pasta_pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pasta_pk)
    projeto = pasta.projeto
    form = PlanilhaForm(request.POST or None)
    if form.is_valid():
        planilha = form.save(commit=False)
        planilha.pasta = pasta
        planilha.save()
        messages.success(request, 'Planilha criada.')
        return redirect('detalhe_planilha', pk=planilha.pk)
    return render(request, 'paineis/form_planilha.html', {
        'form': form, 'pasta': pasta, 'projeto': projeto, 'titulo': 'Nova Planilha',
    })


def detalhe_planilha(request, pk):
    planilha = get_object_or_404(Planilha, pk=pk)
    return render(request, 'paineis/detalhe_planilha.html', {
        'planilha': planilha, 'projeto': planilha.pasta.projeto,
    })


def editar_planilha(request, pk):
    planilha = get_object_or_404(Planilha, pk=pk)
    form = PlanilhaForm(request.POST or None, instance=planilha)
    if form.is_valid():
        form.save()
        messages.success(request, 'Planilha atualizada.')
        return redirect('detalhe_planilha', pk=planilha.pk)
    return render(request, 'paineis/form_planilha.html', {
        'form': form, 'pasta': planilha.pasta, 'projeto': planilha.pasta.projeto,
        'titulo': 'Editar Planilha',
    })


def excluir_planilha(request, pk):
    planilha = get_object_or_404(Planilha, pk=pk)
    projeto_pk = planilha.pasta.projeto.pk
    if request.method == 'POST':
        planilha.delete()
        messages.success(request, 'Planilha excluída.')
        return redirect(_url_projeto(projeto_pk, 'planilhas'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': planilha, 'tipo': 'planilha',
        'voltar_url': 'detalhe_planilha', 'voltar_pk': planilha.pk,
    })


def adicionar_imagem_planilha(request, planilha_pk):
    planilha = get_object_or_404(Planilha, pk=planilha_pk)
    form = ImagemPlanilhaForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        imagem = form.save(commit=False)
        imagem.planilha = planilha
        imagem.save()
        messages.success(request, 'Imagem adicionada.')
        return redirect('detalhe_planilha', pk=planilha.pk)
    return render(request, 'paineis/form_imagem.html', {
        'form': form, 'titulo': 'Adicionar Imagem',
        'voltar_url': 'detalhe_planilha', 'voltar_pk': planilha.pk,
    })


def excluir_imagem_planilha(request, pk):
    imagem = get_object_or_404(ImagemPlanilha, pk=pk)
    planilha_pk = imagem.planilha.pk
    if request.method == 'POST':
        if imagem.imagem and os.path.isfile(imagem.imagem.path):
            os.remove(imagem.imagem.path)
        imagem.delete()
        messages.success(request, 'Imagem excluída.')
        return redirect('detalhe_planilha', pk=planilha_pk)
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': imagem, 'tipo': 'imagem',
        'voltar_url': 'detalhe_planilha', 'voltar_pk': planilha_pk,
    })


# ── Painéis ───────────────────────────────────────────────────────────────────

def criar_painel(request, pasta_pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pasta_pk)
    projeto = pasta.projeto
    form = PainelForm(request.POST or None, pasta=pasta)
    if form.is_valid():
        painel = form.save(commit=False)
        painel.pasta = pasta
        painel.save()
        form.save_m2m()
        messages.success(request, 'Painel criado.')
        return redirect('detalhe_painel', pk=painel.pk)
    return render(request, 'paineis/form_painel.html', {
        'form': form, 'pasta': pasta, 'projeto': projeto, 'titulo': 'Novo Painel',
    })


def detalhe_painel(request, pk):
    painel = get_object_or_404(Painel, pk=pk)
    return render(request, 'paineis/detalhe_painel.html', {
        'painel': painel, 'projeto': painel.pasta.projeto,
    })


def editar_painel(request, pk):
    painel = get_object_or_404(Painel, pk=pk)
    form = PainelForm(request.POST or None, instance=painel, pasta=painel.pasta)
    if form.is_valid():
        form.save()
        messages.success(request, 'Painel atualizado.')
        return redirect('detalhe_painel', pk=painel.pk)
    return render(request, 'paineis/form_painel.html', {
        'form': form, 'pasta': painel.pasta, 'projeto': painel.pasta.projeto,
        'titulo': 'Editar Painel',
    })


def excluir_painel(request, pk):
    painel = get_object_or_404(Painel, pk=pk)
    projeto_pk = painel.pasta.projeto.pk
    if request.method == 'POST':
        painel.delete()
        messages.success(request, 'Painel excluído.')
        return redirect(_url_projeto(projeto_pk, 'paineis'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': painel, 'tipo': 'painel',
        'voltar_url': 'detalhe_painel', 'voltar_pk': painel.pk,
    })


def adicionar_imagem_painel(request, painel_pk):
    painel = get_object_or_404(Painel, pk=painel_pk)
    form = ImagemPainelForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        imagem = form.save(commit=False)
        imagem.painel = painel
        imagem.save()
        messages.success(request, 'Imagem adicionada.')
        return redirect('detalhe_painel', pk=painel.pk)
    return render(request, 'paineis/form_imagem.html', {
        'form': form, 'titulo': 'Adicionar Imagem',
        'voltar_url': 'detalhe_painel', 'voltar_pk': painel.pk,
    })


def excluir_imagem_painel(request, pk):
    imagem = get_object_or_404(ImagemPainel, pk=pk)
    painel_pk = imagem.painel.pk
    if request.method == 'POST':
        if imagem.imagem and os.path.isfile(imagem.imagem.path):
            os.remove(imagem.imagem.path)
        imagem.delete()
        messages.success(request, 'Imagem excluída.')
        return redirect('detalhe_painel', pk=painel_pk)
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': imagem, 'tipo': 'imagem',
        'voltar_url': 'detalhe_painel', 'voltar_pk': painel_pk,
    })


# ── Histórias ─────────────────────────────────────────────────────────────────

def criar_historia(request, pasta_pk):
    pasta = get_object_or_404(PastaDeTrabalho, pk=pasta_pk)
    projeto = pasta.projeto
    form = HistoriaForm(request.POST or None)
    if form.is_valid():
        historia = form.save(commit=False)
        historia.pasta = pasta
        historia.save()
        messages.success(request, 'História criada.')
        return redirect('detalhe_historia', pk=historia.pk)
    return render(request, 'paineis/form_historia.html', {
        'form': form, 'pasta': pasta, 'projeto': projeto, 'titulo': 'Nova História',
    })


def detalhe_historia(request, pk):
    historia = get_object_or_404(Historia, pk=pk)
    return render(request, 'paineis/detalhe_historia.html', {
        'historia': historia, 'projeto': historia.pasta.projeto,
    })


def editar_historia(request, pk):
    historia = get_object_or_404(Historia, pk=pk)
    form = HistoriaForm(request.POST or None, instance=historia)
    if form.is_valid():
        form.save()
        messages.success(request, 'História atualizada.')
        return redirect('detalhe_historia', pk=historia.pk)
    return render(request, 'paineis/form_historia.html', {
        'form': form, 'pasta': historia.pasta, 'projeto': historia.pasta.projeto,
        'titulo': 'Editar História',
    })


def excluir_historia(request, pk):
    historia = get_object_or_404(Historia, pk=pk)
    projeto_pk = historia.pasta.projeto.pk
    if request.method == 'POST':
        historia.delete()
        messages.success(request, 'História excluída.')
        return redirect(_url_projeto(projeto_pk, 'historias'))
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': historia, 'tipo': 'história',
        'voltar_url': 'detalhe_historia', 'voltar_pk': historia.pk,
    })


def adicionar_imagem_historia(request, historia_pk):
    historia = get_object_or_404(Historia, pk=historia_pk)
    form = ImagemHistoriaForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        imagem = form.save(commit=False)
        imagem.historia = historia
        imagem.save()
        messages.success(request, 'Imagem adicionada.')
        return redirect('detalhe_historia', pk=historia.pk)
    return render(request, 'paineis/form_imagem.html', {
        'form': form, 'titulo': 'Adicionar Imagem',
        'voltar_url': 'detalhe_historia', 'voltar_pk': historia.pk,
    })


def excluir_imagem_historia(request, pk):
    imagem = get_object_or_404(ImagemHistoria, pk=pk)
    historia_pk = imagem.historia.pk
    if request.method == 'POST':
        if imagem.imagem and os.path.isfile(imagem.imagem.path):
            os.remove(imagem.imagem.path)
        imagem.delete()
        messages.success(request, 'Imagem excluída.')
        return redirect('detalhe_historia', pk=historia_pk)
    return render(request, 'paineis/confirmar_exclusao.html', {
        'objeto': imagem, 'tipo': 'imagem',
        'voltar_url': 'detalhe_historia', 'voltar_pk': historia_pk,
    })


# ── Exportação Word ───────────────────────────────────────────────────────────

def exportar_word(request, pk):
    projeto = get_object_or_404(Projeto, pk=pk)
    pasta = projeto.pastas.first()
    if not pasta:
        messages.error(request, 'Cadastre os Dados do Tableau antes de exportar.')
        return redirect('detalhe_projeto', pk=pk)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Capa
    titulo_doc = doc.add_heading(pasta.nome, 0)
    titulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Projeto: {projeto.nome}\n').bold = True
    if pasta.modulo_display():
        info.add_run(f'Módulo: {pasta.modulo_display()}\n')
    if pasta.setor_solicitante:
        info.add_run(f'Setor Solicitante: {pasta.setor_solicitante}\n')
    if pasta.solicitante:
        info.add_run(f'Solicitante: {pasta.solicitante}\n')
    from django.utils import timezone
    info.add_run(f'Data: {timezone.localdate().strftime("%d/%m/%Y")}')

    registros = pasta.registros.all()
    if registros.exists():
        doc.add_paragraph()
        tabela_reg = doc.add_table(rows=1, cols=3)
        tabela_reg.style = 'Table Grid'
        h = tabela_reg.rows[0].cells
        h[0].text = 'Ação'
        h[1].text = 'Responsável'
        h[2].text = 'Data'
        for i, reg in enumerate(registros):
            r = tabela_reg.add_row().cells
            r[0].text = 'Desenvolvido por' if i == 0 else 'Atualizado por'
            r[1].text = reg.nome
            r[2].text = reg.data.strftime('%d/%m/%Y')

    # Fontes de Dados
    fontes = pasta.fontes.prefetch_related('filtros', 'imagens').all()
    if fontes.exists():
        doc.add_page_break()
        doc.add_heading('Fontes de Dados', level=1)
        for fonte in fontes:
            doc.add_heading(fonte.nome, level=2)
            tabela_info = doc.add_table(rows=3, cols=2)
            tabela_info.style = 'Table Grid'
            tabela_info.rows[0].cells[0].text = 'Conexão'
            tabela_info.rows[0].cells[1].text = fonte.conexao or '-'
            tabela_info.rows[1].cells[0].text = 'Banco'
            tabela_info.rows[1].cells[1].text = fonte.banco or '-'
            tabela_info.rows[2].cells[0].text = 'Tipo de Conexão'
            tabela_info.rows[2].cells[1].text = fonte.get_tipo_conexao_display()
            doc.add_paragraph()
            filtros = fonte.filtros.all()
            if filtros.exists():
                doc.add_heading('Filtros', level=3)
                tf = doc.add_table(rows=1, cols=2)
                tf.style = 'Table Grid'
                tf.rows[0].cells[0].text = 'Filtro'
                tf.rows[0].cells[1].text = 'Detalhes'
                for f in filtros:
                    r = tf.add_row().cells
                    r[0].text = f.nome
                    r[1].text = f.detalhes or '-'
                doc.add_paragraph()
            if fonte.relacao_conexoes:
                doc.add_heading('Relação de Conexões', level=3)
                doc.add_paragraph(fonte.relacao_conexoes)
            for img in fonte.imagens.all():
                try:
                    doc.add_heading(img.legenda or 'Imagem', level=3)
                    doc.add_picture(img.imagem.path, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass

    # Planilhas
    planilhas = pasta.planilhas.prefetch_related('imagens').all()
    if planilhas.exists():
        doc.add_page_break()
        doc.add_heading('Planilhas', level=1)
        for planilha in planilhas:
            doc.add_heading(planilha.nome, level=2)
            if planilha.descricao:
                doc.add_paragraph(planilha.descricao)
            for img in planilha.imagens.all():
                try:
                    secao = img.secao or ''
                    legenda = img.legenda or 'Imagem'
                    doc.add_heading(f'{secao} — {legenda}' if secao else legenda, level=3)
                    doc.add_picture(img.imagem.path, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass

    # Painéis
    paineis = pasta.paineis.prefetch_related('imagens', 'planilhas').all()
    if paineis.exists():
        doc.add_page_break()
        doc.add_heading('Painéis', level=1)
        for painel in paineis:
            doc.add_heading(painel.nome, level=2)
            if painel.descricao:
                doc.add_paragraph(painel.descricao)
            if painel.regras:
                doc.add_heading('Regras de Negócio', level=3)
                doc.add_paragraph(painel.regras)
            vinculadas = painel.planilhas.all()
            if vinculadas.exists():
                doc.add_heading('Planilhas vinculadas', level=3)
                for pl in vinculadas:
                    doc.add_paragraph(pl.nome, style='List Bullet')
            for img in painel.imagens.all():
                try:
                    secao = img.secao or ''
                    legenda = img.legenda or 'Imagem'
                    doc.add_heading(f'{secao} — {legenda}' if secao else legenda, level=3)
                    doc.add_picture(img.imagem.path, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass

    # Histórias
    historias = pasta.historias.prefetch_related('imagens').all()
    if historias.exists():
        doc.add_page_break()
        doc.add_heading('Histórias', level=1)
        for historia in historias:
            doc.add_heading(historia.nome, level=2)
            if historia.descricao:
                doc.add_paragraph(historia.descricao)
            for img in historia.imagens.all():
                try:
                    secao = img.secao or ''
                    legenda = img.legenda or 'Imagem'
                    doc.add_heading(f'{secao} — {legenda}' if secao else legenda, level=3)
                    doc.add_picture(img.imagem.path, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nome_arquivo = f'{pasta.nome.replace(" ", "_")}_documentacao.docx'
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response
